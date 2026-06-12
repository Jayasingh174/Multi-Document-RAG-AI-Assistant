from docx import Document


def extract_text(path: str) -> str:
    """
    Extract text from a DOCX file including paragraphs and tables.
    Returns a clean text string suitable for RAG ingestion.
    """

    try:
        doc = Document(path)
    except Exception as e:
        raise ValueError(f"Failed to open DOCX file: {e}")

    extracted_text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            extracted_text.append(text)

    # Extract tables with column context
    for table in doc.tables:

        headers = [cell.text.strip() for cell in table.rows[0].cells]

        for row in table.rows[1:]:

            row_data = []

            for i, cell in enumerate(row.cells):

                value = cell.text.strip()

                if value:
                    header = headers[i] if i < len(headers) else f"Column{i}"
                    row_data.append(f"{header}: {value}")

            if row_data:
                extracted_text.append(" | ".join(row_data))

    return "\n".join(extracted_text)